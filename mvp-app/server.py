# -*- coding: utf-8 -*-
"""ParkFlow AI MVP web adapter.

This service does not modify original V0. It reads the existing local data
assets and calls the configured DashScope-compatible LLM to produce reports.
"""
from __future__ import annotations

import io
import json
import mimetypes
import os
import re
import sqlite3
import sys
import threading
import time
import urllib.error
import urllib.parse
import urllib.request
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any


APP_DIR = Path(__file__).resolve().parent
STATIC_DIR = APP_DIR / "static"
WORKSPACE_DIR = APP_DIR.parents[1]
ORIGINAL_DIR = Path(os.getenv("PARKFLOW_ORIGINAL_DIR", str(WORKSPACE_DIR / "original V0"))).resolve()
DB_PATH = ORIGINAL_DIR / "db" / "park_data.db"
GRAPH_PATH = ORIGINAL_DIR / "db" / "industry_graph.json"
RISK_PATH = ORIGINAL_DIR / "external_api" / "risk_data.json"
INTEL_PATH = ORIGINAL_DIR / "external_api" / "company_intelligence.json"
API_KEY_PATH = ORIGINAL_DIR / "api_key.txt"

LLM_API_BASE = os.getenv("DASHSCOPE_API_BASE", "https://dashscope.aliyuncs.com/compatible-mode/v1")
LLM_MODEL = os.getenv("DASHSCOPE_MODEL", "qwen-plus")
LLM_TIMEOUT = int(os.getenv("MVP_LLM_TIMEOUT_SECONDS", "45"))
ALLOWED_ORIGINS = {
    origin.strip().rstrip("/")
    for origin in os.getenv("CORS_ALLOWED_ORIGINS", os.getenv("FRONTEND_ORIGIN", "*")).split(",")
    if origin.strip()
}

INDUSTRY_ALIASES = {
    "半导体": ["半导体", "集成电路", "芯片", "EDA", "晶圆"],
    "人工智能": ["人工智能", "AI", "大模型", "智能算法", "机器学习"],
    "生物医药": ["生物医药", "医药", "医疗", "生命科学", "创新药", "医疗器械", "CRO", "CDMO"],
    "新能源": ["新能源", "储能", "光伏", "电池"],
    "智能制造": ["智能制造", "高端装备", "机器人", "自动化"],
    "低空经济": ["低空经济", "无人机", "航空航天", "eVTOL"],
}

FINANCING_RANK = {
    "未融资": 0,
    "种子轮": 1,
    "天使轮": 2,
    "Pre-A轮": 3,
    "A轮": 4,
    "B轮": 5,
    "C轮": 6,
    "D轮": 7,
    "Pre-IPO": 8,
    "已上市": 9,
}

MATERIAL_LABELS = {
    "outline": "企业拜访提纲",
    "wechat": "微信跟进话术",
    "briefing": "领导汇报摘要",
    "risk": "风险复核清单",
    "plan": "项目推进计划",
    "invite": "招商邀请函",
}

_ORIGINAL_AGENT: Any | None = None
_ORIGINAL_AGENT_LOCK = threading.Lock()


def normalize_thread_context(value: Any) -> dict[str, Any]:
    if not isinstance(value, dict):
        return {}
    return value


def should_use_thread_context(task: str, body: dict[str, Any]) -> bool:
    thread = normalize_thread_context(body.get("thread_context"))
    if not thread:
        return False
    text = task or ""
    contextual_terms = [
        "这家", "该企业", "这家公司", "这个企业", "当前企业", "第1家", "第一家", "第 1 家",
        "上一家", "刚才", "上面", "上述", "它", "他们", "适配吗", "为什么推荐",
        "展开说明", "继续", "基于", "这份", "这批",
    ]
    material_terms = ["邀约", "话术", "拜访", "提纲", "邀请函", "汇报", "材料", "邮件"]
    discovery_terms = ["推荐", "找", "筛选", "找出", "一些", "一批", "哪些", "公司", "企业"]
    industry_terms = ["芯片", "半导体", "医药", "低空", "人工智能", "新能源", "材料", "制造"]
    has_contextual = any(term in text for term in contextual_terms)
    has_material = any(term in text for term in material_terms)
    is_new_discovery = any(term in text for term in discovery_terms) and any(term in text for term in industry_terms)
    has_prior_goal = bool(thread.get("current_goal") or thread.get("history") or thread.get("candidates"))
    vague_followup_discovery = (
        has_prior_goal
        and any(term in text for term in discovery_terms)
        and not any(term in text for term in industry_terms)
        and not re.search(r"[\u4e00-\u9fa5A-Za-z0-9（）()·]{2,40}(?:有限公司|股份有限公司|科技|集团|公司)", text)
    )
    if is_new_discovery and not has_contextual:
        return False
    return has_contextual or has_material or vague_followup_discovery


def is_alternative_recommendation_request(task: str) -> bool:
    text = task or ""
    return bool(
        re.search(r"(还有|其他|其它|换一批|再推荐|继续推荐|别的).*(推荐|企业|公司|名单|线索)", text)
        or re.search(r"这家.*(聊过|看过|排除|不要|不看).*(其他|其它|推荐|企业|公司)", text)
    )


def thread_candidate_names(thread: dict[str, Any]) -> list[str]:
    candidates = thread.get("candidates") if isinstance(thread.get("candidates"), list) else []
    names: list[str] = []
    for item in candidates:
        if not isinstance(item, dict):
            continue
        name = str(item.get("name") or "").strip()
        if name and name not in names:
            names.append(name)
    return names


def excluded_names_from_task(task: str) -> list[str]:
    names: list[str] = []
    for match in re.findall(r"[\u4e00-\u9fa5A-Za-z0-9（）()·]{2,40}(?:有限公司|股份有限公司|科技|集团|公司)", task or ""):
        clean = match.strip(" ，,;；。:：")
        if clean and clean not in names:
            names.append(clean)
    return names


def contextualize_task_for_backend(task: str, body: dict[str, Any]) -> str:
    """Attach current frontend thread context without changing backend workflow."""
    thread = normalize_thread_context(body.get("thread_context"))
    if not thread or not should_use_thread_context(task, body):
        return task

    parts = [f"用户当前追问：{task}"]
    original_goal = str(thread.get("current_goal") or "").strip()
    alt_recommendation = is_alternative_recommendation_request(task)
    active_company = "" if alt_recommendation else str(thread.get("active_company") or "").strip()
    last_summary = str(thread.get("last_summary") or "").strip()

    if alt_recommendation and original_goal:
        names = thread_candidate_names(thread)
        exclude = "、".join(names[:12])
        task = f"{original_goal}\n补充要求：继续推荐其他企业，不要再围绕已讨论企业展开。"
        if exclude:
            task += f"\n排除已推荐企业：{exclude}"
        task += "\n请直接给出新的候选企业、推荐理由、风险提示和下一步推进建议。"
        return task

    if original_goal:
        parts.append(f"当前招商目标：{original_goal[:500]}")
    if active_company:
        parts.append(f"当前关注企业：{active_company[:120]}")
    if last_summary:
        parts.append(f"上一轮结论摘要：{last_summary[:700]}")

    candidates = thread.get("candidates") if isinstance(thread.get("candidates"), list) else []
    if candidates:
        lines = []
        for item in candidates[:10]:
            if not isinstance(item, dict):
                continue
            name = str(item.get("name") or "").strip()
            score = item.get("score", "")
            industry = str(item.get("industry") or "").strip()
            reason = str(item.get("reason") or "").strip()
            if name:
                lines.append(f"- {name}｜匹配分 {score}｜{industry}｜{reason[:160]}")
        if lines:
            parts.append("上一轮推荐企业：\n" + "\n".join(lines))

    history = thread.get("history") if isinstance(thread.get("history"), list) else []
    if history:
        lines = []
        for item in history[-6:]:
            if not isinstance(item, dict):
                continue
            role = str(item.get("role") or "").strip()
            content = str(item.get("content") or "").strip()
            if role and content:
                lines.append(f"{role}: {content[:220]}")
        if lines:
            parts.append("最近对话：\n" + "\n".join(lines))

    return "\n\n".join(parts) + "\n\n请基于以上同一招商线程上下文回答，不要把“这家企业”“第1家”等追问理解成新任务。"


def should_use_original_agent_chat(task: str) -> bool:
    """Route non-project conversation to the original Agent instead of forcing analysis."""
    text = (task or "").strip()
    if not text:
        return False
    meta_patterns = [
        r"你是谁",
        r"你能做什么",
        r"介绍一下",
        r"怎么使用",
        r"如何使用",
        r"帮助",
        r"help",
        r"你好",
    ]
    if any(re.search(pattern, text, re.IGNORECASE) for pattern in meta_patterns):
        return True
    business_signals = [
        "招商", "企业", "公司", "产业", "政策", "风险", "园区", "入驻", "推荐", "筛选", "研判",
        "分析", "评估", "生成", "话术", "邮件", "邀请函", "拜访", "汇报", "租金", "营收",
    ]
    return len(text) <= 36 and not any(signal in text for signal in business_signals)


def should_start_structured_workflow(task: str, body: dict[str, Any] | None = None) -> bool:
    """Only start the structured investment workflow for explicit business tasks.

    Ambiguous delegation such as "可以帮我找到合适的企业吗" should stay conversational so
    the model can clarify goals instead of the adapter inventing filters.
    """
    text = (task or "").strip()
    if not text:
        return False
    body = body or {}
    count_requested = requested_count(text, 0) > 0
    industry_requested = bool(requested_industry_terms(text)[1])
    discovery_words = ["推荐", "找出", "找到", "筛选", "寻找", "匹配", "调研", "盘点", "候选", "线索", "重点推进", "找一下"]
    enterprise_words = ["企业", "公司", "对象", "线索", "候选", "名单"]
    outcome_words = ["推荐理由", "风险提示", "下一步", "推进建议", "政策抓手", "适配", "合适", "招商建议"]
    capability_question = bool(re.search(r"(你可以|你能|能不能|可以帮我|能帮我).*(吗|么|？|\?)", text))
    has_discovery = any(word in text for word in discovery_words)
    has_enterprise = any(word in text for word in enterprise_words)
    has_outcome = any(word in text for word in outcome_words)

    if capability_question and not (count_requested or industry_requested or has_outcome):
        return False
    if detect_intent(text) == "data_inventory":
        return True
    if infer_company_name(text):
        return True
    if is_alternative_recommendation_request(text):
        return True
    thread = normalize_thread_context(body.get("thread_context"))
    if thread and has_discovery and count_requested and not industry_requested:
        return True
    if has_discovery and (has_enterprise or industry_requested or count_requested or has_outcome):
        return True
    if has_enterprise and count_requested and any(word in text for word in ["适合", "合适", "重点", "推进", "招商"]):
        return True
    if industry_requested and any(word in text for word in ["分析", "评估", "研判", "重点推进"]):
        return True
    if any(word in text for word in ["产业链", "产业机会", "产业分析", "政策", "抓手", "补贴"]) and industry_requested:
        return True
    if any(word in text for word in ["生成", "写", "起草", "邀请函", "话术", "拜访提纲", "汇报材料"]):
        thread = normalize_thread_context(body.get("thread_context"))
        return bool(thread.get("current_goal") or thread.get("active_company") or thread.get("candidates") or body.get("company"))
    return False


def get_original_agent_unlocked() -> Any:
    global _ORIGINAL_AGENT
    if _ORIGINAL_AGENT is not None:
        return _ORIGINAL_AGENT
    if str(ORIGINAL_DIR) not in sys.path:
        sys.path.insert(0, str(ORIGINAL_DIR))
    os.environ.setdefault("ENABLE_PLAYWRIGHT_MCP", "0")
    old_cwd = os.getcwd()
    try:
        os.chdir(ORIGINAL_DIR)
        from agent.agent_loop import RecruitmentAgent  # type: ignore

        _ORIGINAL_AGENT = RecruitmentAgent()
        return _ORIGINAL_AGENT
    finally:
        os.chdir(old_cwd)


def build_original_style_prompt() -> str:
    stats = data_stats()
    return (
        "你是 ParkFlow，面向产业园区招商工作的智能顾问。"
        "你的回答应像专业招商主管或产业研究顾问，清楚、克制、可继续追问。"
        "普通能力说明或闲聊只需自然回答，不要假装已经完成企业检索；"
        "用户提出明确招商任务时，简短确认会进入企业筛选、研判或材料生成流程。"
        "禁止暴露内部实现词：OPC、RAG、MCP、Workflow、tool_code、get_company_risk、ChromaDB、"
        "requested_industry、candidate_enterprises、park_policies、industry_graph、JSON字段。"
        "不要使用乱码符号或 emoji。格式使用短段落和简洁列表。"
        f"当前可用本地资料：{stats.get('enterprise_count')} 家企业线索、"
        f"{stats.get('policy_count')} 条政策资料、{stats.get('resource_count')} 项园区资源。"
    )


def call_original_style_chat(task: str) -> str:
    system = build_original_style_prompt()
    return call_llm(
        [
            {"role": "system", "content": system},
            {"role": "user", "content": task},
        ],
        max_tokens=1400,
    )


def read_json(path: Path, fallback: Any) -> Any:
    try:
        with path.open("r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return fallback


def get_api_key() -> str | None:
    env_key = os.getenv("DASHSCOPE_API_KEY")
    if env_key:
        return env_key.strip()
    if API_KEY_PATH.exists():
        key = API_KEY_PATH.read_text(encoding="utf-8").strip()
        return key or None
    return None


def db_rows(sql: str, params: tuple[Any, ...] = ()) -> list[dict[str, Any]]:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    try:
        return [dict(row) for row in conn.execute(sql, params).fetchall()]
    finally:
        conn.close()


def db_one(sql: str, params: tuple[Any, ...] = ()) -> dict[str, Any] | None:
    rows = db_rows(sql, params)
    return rows[0] if rows else None


def all_enterprises() -> list[dict[str, Any]]:
    return db_rows("SELECT * FROM enterprises ORDER BY id LIMIT 500")


def data_stats() -> dict[str, Any]:
    enterprise_count = db_one("SELECT COUNT(*) AS count FROM enterprises") or {"count": 0}
    industry_count = db_one("SELECT COUNT(DISTINCT industry) AS count FROM enterprises") or {"count": 0}
    tenant_count = db_one("SELECT COUNT(*) AS count FROM crm_records WHERE stage='已签约'") or {"count": 0}
    policy_count = db_one("SELECT COUNT(*) AS count FROM park_policies") or {"count": 0}
    resource_count = db_one("SELECT COUNT(*) AS count FROM park_resources") or {"count": 0}
    crm_count = db_one("SELECT COUNT(*) AS count FROM crm_records") or {"count": 0}
    top_industries = db_rows(
        "SELECT industry, COUNT(*) AS count FROM enterprises "
        "GROUP BY industry ORDER BY count DESC, industry LIMIT 8"
    )
    return {
        "enterprise_count": enterprise_count.get("count", 0),
        "tenant_count": tenant_count.get("count", 0),
        "industry_count": industry_count.get("count", 0),
        "policy_count": policy_count.get("count", 0),
        "resource_count": resource_count.get("count", 0),
        "crm_count": crm_count.get("count", 0),
        "top_industries": top_industries,
    }


def requested_industry_terms(task: str) -> tuple[str, list[str]]:
    task = task or ""
    for canonical, aliases in INDUSTRY_ALIASES.items():
        if any(alias.lower() in task.lower() for alias in aliases):
            return canonical, aliases

    for row in db_rows("SELECT DISTINCT industry FROM enterprises WHERE industry IS NOT NULL AND industry != ''"):
        industry = str(row.get("industry") or "")
        roots = [industry, industry.split("/")[0], industry.split("（")[0]]
        if any(root and root in task for root in roots):
            return industry, roots
    return "", []


def detect_intent(task: str) -> str:
    text = task or ""
    if re.search(r"(收录|多少|几家|数量|数据库|企业池).*(企业|公司|产业|政策|资源)|企业.*(多少|几家|收录|数量)", text):
        return "data_inventory"
    if any(word in text for word in ["生成", "写", "起草", "邀请函", "话术", "拜访提纲", "汇报材料"]):
        return "material_generation"
    if is_industry_recommendation_task(text):
        return "industry_recommendation"
    if any(word in text for word in ["推荐", "找出", "筛选", "寻找", "匹配", "调研", "盘点", "候选", "线索"]) and requested_industry_terms(text)[1]:
        return "mission_discovery"
    if any(word in text for word in ["产业链", "产业机会", "产业分析", "缺口"]):
        return "industry_analysis"
    if any(word in text for word in ["政策", "抓手", "补贴"]):
        return "policy_analysis"
    if infer_company_name(text):
        return "company_analysis"
    return "mission_discovery"


def is_industry_recommendation_task(task: str) -> bool:
    if not task:
        return False
    has_direction = any(word in task for word in ["方向", "产业", "行业", "赛道", "领域"])
    has_recommend = any(word in task for word in ["找出", "推荐", "最值得", "优先", "重点推进", "政策抓手"])
    return has_direction and has_recommend and bool(requested_industry_terms(task)[1])


def requested_count(task: str, default: int = 5) -> int:
    text = task or ""
    match = re.search(r"(\d{1,2})\s*[家个名]", text)
    if match:
        return max(1, min(12, int(match.group(1))))
    chinese_numbers = {
        "一": 1,
        "二": 2,
        "两": 2,
        "三": 3,
        "四": 4,
        "五": 5,
        "六": 6,
        "七": 7,
        "八": 8,
        "九": 9,
        "十": 10,
    }
    match = re.search(r"([一二两三四五六七八九十])\s*[家个名]", text)
    if match:
        return chinese_numbers.get(match.group(1), default)
    return default


def bounded_score(value: float) -> int:
    return max(0, min(100, int(round(value))))


def revenue_score(value: str) -> int:
    value = value or ""
    if "10亿以上" in value:
        return 100
    if "5亿-10亿" in value:
        return 92
    if "3亿-5亿" in value:
        return 84
    if "1亿-3亿" in value:
        return 76
    if "5000万-1亿" in value:
        return 66
    if "2000万-5000万" in value:
        return 54
    if "500万-2000万" in value:
        return 42
    if "0-500万" in value:
        return 28
    if "未盈利" in value:
        return 22
    return 45


def parse_tags(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item) for item in value]
    if not value:
        return []
    try:
        parsed = json.loads(str(value))
        if isinstance(parsed, list):
            return [str(item) for item in parsed]
    except Exception:
        pass
    return [item.strip() for item in re.split(r"[,，、\s]+", str(value)) if item.strip()]


def mission_weights(task: str) -> dict[str, float]:
    text = task or ""
    weights = {"contribution": 0.34, "rent": 0.22, "growth": 0.28, "risk": 0.16}
    if any(word in text for word in ["贡献", "税收", "产值", "就业", "产业链", "补链", "强链"]):
        weights["contribution"] += 0.12
    if any(word in text for word in ["租金", "交得起", "承租", "现金流", "收入"]):
        weights["rent"] += 0.14
    if any(word in text for word in ["发展", "成长", "潜力", "融资", "高成长", "未来"]):
        weights["growth"] += 0.12
    if any(word in text for word in ["风险", "稳健", "可靠", "合规"]):
        weights["risk"] += 0.08
    total = sum(weights.values()) or 1
    return {key: value / total for key, value in weights.items()}


def mission_score(row: dict[str, Any], task: str, canonical: str = "", terms: list[str] | None = None) -> dict[str, Any]:
    terms = terms or []
    employees = int(row.get("employees") or 0)
    patents = int(row.get("patents") or 0)
    financing = str(row.get("financing_stage") or "")
    financing_value = max((rank for key, rank in FINANCING_RANK.items() if key in financing), default=1)
    tags = parse_tags(row.get("tags"))
    risk = risk_for(str(row.get("name") or ""), row)
    risk_value = float(risk.get("risk_score") or 35)
    industry = str(row.get("industry") or "")
    sub_industry = str(row.get("sub_industry") or "")
    haystack = f"{industry} {sub_industry} {' '.join(tags)} {row.get('brief') or ''}"

    industry_match = 0
    if canonical and (canonical in industry or canonical in sub_industry):
        industry_match = 24
    elif terms and any(term and term in haystack for term in terms):
        industry_match = 16
    elif any(key in industry for key in ["半导体", "人工智能", "生物医药", "新能源", "智能制造", "低空经济"]):
        industry_match = 12

    tag_bonus = 0
    for tag in tags:
        if any(word in tag for word in ["专精特新", "瞪羚", "高新", "小巨人", "独角兽", "上市公司子公司"]):
            tag_bonus += 4
    contribution = bounded_score(28 + industry_match + min(24, employees / 10) + min(20, patents / 6) + min(12, tag_bonus))
    rent = bounded_score(revenue_score(str(row.get("revenue_range") or "")) * 0.72 + min(18, employees / 18) + financing_value * 2.2)
    growth = bounded_score(24 + financing_value * 6 + min(24, patents / 5) + min(18, tag_bonus) + (8 if int(row.get("founded_year") or 0) >= 2018 else 0))
    risk_score = bounded_score(100 - risk_value)
    weights = mission_weights(task)
    total = bounded_score(
        contribution * weights["contribution"]
        + rent * weights["rent"]
        + growth * weights["growth"]
        + risk_score * weights["risk"]
    )
    reasons = []
    if industry_match:
        reasons.append("产业方向与园区重点赛道或任务语义匹配")
    if revenue_score(str(row.get("revenue_range") or "")) >= 60:
        reasons.append(f"收入区间 {row.get('revenue_range')}，具备较好租金承载力")
    if patents >= 40:
        reasons.append(f"专利 {patents} 项，技术积累较强")
    if tags:
        reasons.append("具备 " + "、".join(tags[:3]) + " 等标签")
    if not reasons:
        reasons.append("基础画像完整，可作为初筛线索进入人工复核")
    return {
        "total": total,
        "contribution": contribution,
        "rent": rent,
        "growth": growth,
        "risk": risk_score,
        "reasons": reasons[:3],
    }


def mission_candidates(task: str, limit: int | None = None) -> tuple[str, list[dict[str, Any]]]:
    limit = limit or requested_count(task)
    canonical, terms = requested_industry_terms(task)
    excluded = set(excluded_names_from_task(task))
    rows = []
    for row in all_enterprises():
        if row.get("name") in excluded:
            continue
        haystack = f"{row.get('industry', '')} {row.get('sub_industry', '')} {row.get('tags', '')} {row.get('brief', '')}"
        if terms and not any(term and term in haystack for term in terms):
            continue
        enriched = dict(row)
        scores = mission_score(enriched, task, canonical, terms)
        enriched["_score"] = scores["total"]
        enriched["_score_parts"] = {
            "贡献度": scores["contribution"],
            "租金承载": scores["rent"],
            "成长性": scores["growth"],
            "风险安全": scores["risk"],
        }
        enriched["_rank_reason"] = "；".join(scores["reasons"])
        rows.append(enriched)
    rows.sort(key=lambda item: item.get("_score", 0), reverse=True)
    return canonical, rows[:limit]


def enterprise_score(row: dict[str, Any], canonical: str = "", terms: list[str] | None = None) -> float:
    employees = int(row.get("employees") or 0)
    patents = int(row.get("patents") or 0)
    risk = row.get("risk_score")
    risk_value = float(risk) if isinstance(risk, (int, float)) and risk else 30.0
    financing = str(row.get("financing_stage") or "")
    financing_value = max((rank for key, rank in FINANCING_RANK.items() if key in financing), default=1)
    revenue_bonus = 10 if "盈利" in str(row.get("revenue_range") or "") else 0
    industry = str(row.get("industry") or "")
    sub_industry = str(row.get("sub_industry") or "")
    tags = str(row.get("tags") or "")
    terms = terms or []
    match_bonus = 0
    if canonical and canonical in industry:
        match_bonus += 420
    elif any(term and term in industry for term in terms):
        match_bonus += 110
    if canonical and canonical in sub_industry:
        match_bonus += 90
    elif any(term and term in sub_industry for term in terms):
        match_bonus += 55
    if any(term and term in tags for term in terms):
        match_bonus += 20
    return match_bonus + employees * 0.25 + patents * 3 + financing_value * 8 + revenue_bonus - risk_value * 0.45


def industry_candidates(task: str) -> tuple[str, list[dict[str, Any]]]:
    canonical, terms = requested_industry_terms(task)
    if not terms:
        return "", []
    rows = []
    for row in all_enterprises():
        haystack = f"{row.get('industry', '')} {row.get('sub_industry', '')} {row.get('tags', '')}"
        if any(term and term in haystack for term in terms):
            enriched = dict(row)
            enriched["_score"] = round(enterprise_score(enriched, canonical, terms), 2)
            rows.append(enriched)
    rows.sort(key=lambda item: item.get("_score", 0), reverse=True)
    return canonical, rows[:8]


def infer_company_name(task: str, explicit: str | None = None) -> str | None:
    _PLACEHOLDER_RE = re.compile(r"[【〔﹝［<][^】〕﹞］>]{1,30}[】〕﹞］>]")
    if explicit:
        explicit = explicit.strip()
        if _PLACEHOLDER_RE.search(explicit):
            return None
        return explicit
    task = task or ""
    for company in all_enterprises():
        name = str(company.get("name") or "")
        short = name.replace("有限公司", "").replace("科技", "")
        if name and name in task:
            return name
        if short and len(short) >= 3 and short in task:
            return name
    match = re.search(r"分析(.+?)(?:是否|值不值得|适不适合|能否|并|，|。|$)", task)
    if match:
        name = match.group(1).strip("“”\"' ")
        if _PLACEHOLDER_RE.search(name):
            return None
        return name
    return None


def find_enterprise(company_name: str | None, task: str = "") -> dict[str, Any] | None:
    if company_name:
        exact = db_one("SELECT * FROM enterprises WHERE name = ? LIMIT 1", (company_name,))
        if exact:
            return exact
        like = db_one("SELECT * FROM enterprises WHERE name LIKE ? LIMIT 1", (f"%{company_name}%",))
        if like:
            return like

    canonical, terms = requested_industry_terms(task)
    if terms:
        _, candidates = industry_candidates(task or canonical)
        return candidates[0] if candidates else None
    return None


def risk_for(company_name: str, enterprise: dict[str, Any] | None) -> dict[str, Any]:
    risk_data = read_json(RISK_PATH, {})
    for key, value in risk_data.items():
        if company_name in key or key in company_name:
            return {"found": True, "source": "risk_data.json", **value}
    if not enterprise:
        return {"found": False, "message": "未找到企业风险数据"}

    risk_score = enterprise.get("risk_score")
    if not isinstance(risk_score, (int, float)) or risk_score == 0:
        employees = int(enterprise.get("employees") or 0)
        patents = int(enterprise.get("patents") or 0)
        score = 30
        if employees < 30:
            score += 12
        if employees > 100:
            score -= 8
        if patents >= 20:
            score -= 8
        if "未盈利" in str(enterprise.get("revenue_range", "")):
            score += 10
        risk_score = max(5, min(85, score))

    if risk_score < 30:
        level = "低风险"
    elif risk_score < 60:
        level = "中等风险"
    else:
        level = "高风险"
    return {
        "found": True,
        "source": "enterprises.risk_score + 本地规则",
        "risk_score": risk_score,
        "risk_level": level,
        "risk_tags": [],
        "operating_status": "正常",
    }


def crm_for(company_name: str) -> list[dict[str, Any]]:
    return db_rows(
        "SELECT * FROM crm_records WHERE company_name LIKE ? ORDER BY last_contact DESC",
        (f"%{company_name}%",),
    )


def policies_for(industry: str, extra_terms: list[str] | None = None) -> list[dict[str, Any]]:
    root = industry.split("/")[0] if industry else ""
    terms = [industry, root, *(extra_terms or [])]
    terms = [term for term in dict.fromkeys(terms) if term]
    rows: list[dict[str, Any]] = []
    for term in terms:
        rows.extend(
            db_rows(
                "SELECT * FROM park_policies WHERE target LIKE ? OR name LIKE ? OR condition LIKE ? LIMIT 8",
                (f"%{term}%", f"%{term}%", f"%{term}%"),
            )
        )
    seen = set()
    rows = [row for row in rows if not (row.get("id") in seen or seen.add(row.get("id")))]
    if rows:
        return rows[:8]
    return db_rows("SELECT * FROM park_policies LIMIT 8")


def resources_for() -> list[dict[str, Any]]:
    return db_rows(
        "SELECT * FROM park_resources WHERE status = '空置' "
        "ORDER BY rent_per_sqm ASC, area_sqm DESC LIMIT 6"
    )


def intelligence_for(company_name: str) -> dict[str, Any]:
    data = read_json(INTEL_PATH, {})
    for key, value in data.items():
        if company_name in key or key in company_name or company_name.replace("有限公司", "") in key:
            return {"found": True, "source": "company_intelligence.json", **value}
    return {"found": False, "message": "未找到外部情报"}


def industry_context(industry: str) -> dict[str, Any]:
    graph = read_json(GRAPH_PATH, {})
    target = None
    for node in graph.get("nodes", []):
        node_id = str(node.get("id") or "")
        if industry in node_id or node_id in industry or industry.split("/")[0] in node_id:
            target = node_id
            break
    if not target:
        return {"found": False, "message": f"未找到 {industry} 的产业图谱"}

    upstream = []
    downstream = []
    for link in graph.get("links", []):
        if link.get("target") == target:
            upstream.append({"name": link.get("source"), "relation": link.get("relation")})
        if link.get("source") == target:
            downstream.append({"name": link.get("target"), "relation": link.get("relation")})
    gaps = graph.get("graph", {}).get("chain_gaps", {})
    gap = {}
    for key, value in gaps.items():
        if key in target or target in key or key in industry:
            gap = value
            break
    return {
        "found": True,
        "industry": target,
        "upstream": upstream,
        "downstream": downstream,
        "gap_analysis": gap,
        "source": "industry_graph.json",
    }


def rag_for(task: str, enterprise: dict[str, Any] | None, industry: str) -> list[dict[str, Any]]:
    """Best-effort real RAG retrieval from original V0 ChromaDB."""
    try:
        sys.path.insert(0, str(ORIGINAL_DIR))
        from data.init_chromadb import get_vector_store  # type: ignore

        store = get_vector_store()
        query_parts = [task]
        if enterprise:
            query_parts.append(str(enterprise.get("name") or ""))
            query_parts.append(str(enterprise.get("brief") or ""))
        if industry:
            query_parts.append(industry)
        results = store.query(query=" ".join(query_parts), collection="all", top_k=4)
        return [
            {
                "id": item.get("id", ""),
                "collection": item.get("collection", ""),
                "content": item.get("content", ""),
                "metadata": item.get("metadata") or {},
                "similarity": item.get("similarity"),
            }
            for item in results
        ]
    except (ImportError, ModuleNotFoundError) as exc:
        return [
            {
                "id": "rag_deps_missing",
                "collection": "rag",
                "content": "向量知识库依赖未安装（chromadb / openai），已使用结构化企业库、政策库、CRM、产业图谱和风险数据完成本次研判。",
                "metadata": {"error": str(exc), "reason": "deps_missing"},
                "unavailable": True,
            }
        ]
    except ValueError as exc:
        return [
            {
                "id": "rag_no_api_key",
                "collection": "rag",
                "content": "向量知识库的 Embedding API Key 未配置，已使用结构化企业库、政策库、CRM、产业图谱和风险数据完成本次研判。",
                "metadata": {"error": str(exc), "reason": "no_api_key"},
                "unavailable": True,
            }
        ]
    except Exception as exc:
        return [
            {
                "id": "rag_unavailable",
                "collection": "rag",
                "content": f"向量知识库暂时不可用（{str(exc)[:120]}），已使用结构化企业库、政策库、CRM、产业图谱和风险数据完成本次研判。",
                "metadata": {"error": str(exc), "reason": "unknown"},
                "unavailable": True,
            }
        ]


_PLACEHOLDER_PATTERN = re.compile(r"[【〔﹝［<][^】〕﹞］>]{1,30}[】〕﹞］>]")

def _has_placeholder(task: str) -> bool:
    return bool(_PLACEHOLDER_PATTERN.search(task or ""))


def collect_context(task: str, company_name: str | None = None) -> dict[str, Any]:
    intent = detect_intent(task)
    requested_industry = ""
    target_count = requested_count(task)
    candidates: list[dict[str, Any]] = []

    if _has_placeholder(task):
        return {
            "ok": False,
            "intent": intent,
            "error": "检测到未填写的待补充内容（如【企业名称】），请先将【……】替换为具体信息后再开始分析。",
            "sources": [],
        }

    if intent in ("industry_recommendation", "mission_discovery"):
        requested_industry, candidates = mission_candidates(task, target_count)
        enterprise = candidates[0] if candidates else None
    else:
        explicit = infer_company_name(task, company_name)
        if not explicit and requested_industry_terms(task)[1]:
            requested_industry, candidates = mission_candidates(task, target_count)
            enterprise = candidates[0] if candidates else None
        else:
            enterprise = find_enterprise(explicit, task)
        if not enterprise:
            if _has_placeholder(task):
                return {
                    "ok": False,
                    "intent": intent,
                    "error": "检测到未填写的待补充内容（如【企业名称】），请先将【……】替换为具体信息后再开始分析。",
                    "sources": [],
                }
            requested_industry, candidates = mission_candidates(task, target_count)
            enterprise = candidates[0] if candidates else None

    if not enterprise:
        return {
            "ok": False,
            "intent": intent,
            "error": "Agent 已读取本地企业库，但没有找到足够匹配的线索。可以补充产业方向、企业特征或招商目标后继续。",
            "sources": [{"id": "enterprises", "name": "本地企业库", "status": "not_found"}],
        }

    name = str(enterprise["name"])
    industry = str(enterprise.get("industry") or "")
    industry_info = industry_context(industry)
    policies = policies_for(industry, [requested_industry] if requested_industry else [])
    context = {
        "ok": True,
        "task": task,
        "intent": intent,
        "company_name": name,
        "enterprise": enterprise,
        "risk": risk_for(name, enterprise),
        "crm_records": crm_for(name),
        "policies": policies,
        "resources": resources_for(),
        "industry": industry_info,
        "external_intelligence": intelligence_for(name),
        "rag_results": rag_for(task, enterprise, industry),
        "requested_industry": requested_industry,
        "requested_count": target_count,
        "candidate_enterprises": candidates,
        "discovery_mode": intent in ("industry_recommendation", "mission_discovery") or bool(candidates),
    }
    context["sources"] = build_sources(context)
    context["evidence"] = build_evidence(context)
    context["steps"] = build_steps(context)
    context["actions"] = build_actions(context)
    return context


def build_sources(context: dict[str, Any]) -> list[dict[str, Any]]:
    enterprise = context.get("enterprise") or {}
    sources = [
        {
            "id": "enterprises",
            "name": "本地企业库",
            "status": "hit",
            "detail": enterprise.get("name", ""),
        }
    ]
    if context.get("candidate_enterprises"):
        sources.append({"id": "candidate_pool", "name": "候选企业池", "status": "hit", "detail": f"{len(context['candidate_enterprises'])} 家"})
    if context.get("crm_records"):
        sources.append({"id": "crm_records", "name": "CRM 跟进记录", "status": "hit", "detail": f"{len(context['crm_records'])} 条"})
    if context.get("policies"):
        sources.append({"id": "park_policies", "name": "园区政策库", "status": "hit", "detail": f"{len(context['policies'])} 条"})
    if context.get("resources"):
        sources.append({"id": "park_resources", "name": "园区资源库", "status": "hit", "detail": f"{len(context['resources'])} 条"})
    if context.get("industry", {}).get("found"):
        sources.append({"id": "industry_graph", "name": "产业图谱", "status": "hit", "detail": context["industry"].get("industry", "")})
    if context.get("external_intelligence", {}).get("found"):
        sources.append({"id": "external_intel", "name": "本地外部情报", "status": "hit", "detail": "工商/专利/融资"})
    if context.get("risk", {}).get("found"):
        sources.append({"id": "risk_data", "name": "风险数据", "status": "hit", "detail": context["risk"].get("risk_level", "")})
    rag_results = context.get("rag_results") or []
    if rag_results:
        if all(item.get("unavailable") for item in rag_results):
            reason = (rag_results[0].get("metadata") or {}).get("reason", "")
            if reason == "deps_missing":
                detail = "依赖未安装"
            elif reason == "no_api_key":
                detail = "Embedding Key 未配置"
            else:
                detail = "暂时不可用"
            sources.append({"id": "rag", "name": "向量知识库", "status": "unavailable", "detail": detail})
        else:
            sources.append({"id": "rag", "name": "向量知识库", "status": "hit", "detail": f"{len(rag_results)} 条片段"})
    return sources


def build_evidence(context: dict[str, Any]) -> list[dict[str, Any]]:
    enterprise = context.get("enterprise") or {}
    evidence = [
        {
            "id": "enterprise_profile",
            "title": "企业画像",
            "source": "本地企业库 enterprises",
            "type": "structured_db",
            "snippet": f"{enterprise.get('name')}，{enterprise.get('industry')} / {enterprise.get('sub_industry')}，员工 {enterprise.get('employees')} 人，融资阶段 {enterprise.get('financing_stage')}，专利 {enterprise.get('patents')} 项。",
            "match_reason": "用于判断企业基础实力、产业方向和成长阶段。",
            "relation": "支撑核心匹配度和推进优先级。",
        }
    ]

    if context.get("candidate_enterprises"):
        top = context["candidate_enterprises"][: context.get("requested_count", 5)]
        evidence.append(
            {
                "id": "candidate_summary",
                "title": "推荐依据摘要",
                "source": "本地企业画像与任务约束",
                "type": "ranking",
                "snippet": f"按当前任务要求筛选出 {len(top)} 家候选企业：" + "；".join([f"{item.get('name')}（{item.get('industry')}）" for item in top]),
                "match_reason": "综合产业方向、营收与租金承载、成长性、风险安全和政策适配进行初筛。",
                "relation": "用于解释为什么这些企业进入本次推荐名单。",
            }
        )

    crm_records = context.get("crm_records") or []
    if crm_records:
        crm = crm_records[0]
        evidence.append(
            {
                "id": "crm_records",
                "title": "历史触达",
                "source": "CRM 跟进记录",
                "type": "crm",
                "snippet": f"{crm.get('stage')}，意向 {crm.get('interest_level')}，最近联系 {crm.get('last_contact')}：{crm.get('notes')}",
                "match_reason": "用于判断是否已有沟通基础和下一步推进节奏。",
                "relation": "支撑行动建议和材料语气。",
            }
        )

    for policy in (context.get("policies") or [])[:3]:
        evidence.append(
            {
                "id": f"policy_{policy.get('id')}",
                "title": policy.get("name", "政策依据"),
                "source": "园区政策库 park_policies",
                "type": "policy",
                "snippet": f"{policy.get('type')}：{policy.get('amount')}。适用对象：{policy.get('target')}。条件：{policy.get('condition')}",
                "match_reason": "用于判断招商政策抓手是否可落地。",
                "relation": "支撑政策匹配和材料生成。",
            }
        )

    industry = context.get("industry") or {}
    if industry.get("found"):
        upstream = "、".join([item.get("name", "") for item in industry.get("upstream", [])[:3]]) or "暂无上游片段"
        downstream = "、".join([item.get("name", "") for item in industry.get("downstream", [])[:3]]) or "暂无下游片段"
        evidence.append(
            {
                "id": "industry_graph",
                "title": "产业链位置",
                "source": "产业图谱 industry_graph.json",
                "type": "graph",
                "snippet": f"节点：{industry.get('industry')}。上游：{upstream}。下游：{downstream}。",
                "match_reason": "用于判断企业是否补强园区产业链。",
                "relation": "支撑产业互补性和产业招商逻辑。",
            }
        )

    risk = context.get("risk") or {}
    if risk.get("found"):
        evidence.append(
            {
                "id": "risk_data",
                "title": "风险核验",
                "source": risk.get("source", "风险数据"),
                "type": "risk",
                "snippet": f"风险评分 {risk.get('risk_score')}，等级 {risk.get('risk_level')}，经营状态 {risk.get('operating_status', '未注明')}。",
                "match_reason": "用于判断是否进入重点推进或需先复核。",
                "relation": "支撑风险提示和推进节奏。",
            }
        )

    for item in (context.get("rag_results") or [])[:3]:
        content = str(item.get("content") or "")
        if item.get("unavailable"):
            evidence.append(
                {
                    "id": "rag_unavailable",
                    "title": "向量知识库状态",
                    "source": "RAG / ChromaDB",
                    "type": "system_status",
                    "snippet": content[:220],
                    "match_reason": "运行时依赖未连接，系统没有把该项作为结论依据。",
                    "relation": "提示当前链路缺口，结论仍基于已命中的结构化数据源。",
                }
            )
        else:
            evidence.append(
                {
                    "id": item.get("id") or "rag",
                    "title": (item.get("metadata") or {}).get("title") or "知识库片段",
                    "source": f"向量知识库 {item.get('collection') or ''}".strip(),
                    "type": "rag",
                    "snippet": content[:220],
                    "match_reason": "通过语义检索补充政策、案例或产业知识。",
                    "relation": "用于增强报告依据，但不替代结构化数据库。",
                }
            )
    return evidence


def build_steps(context: dict[str, Any]) -> list[dict[str, Any]]:
    if context.get("discovery_mode"):
        db_detail = f"已筛出 {len(context.get('candidate_enterprises') or [])} 家候选企业"
    else:
        db_detail = context.get("company_name", "")
    return [
        {"id": "intent", "label": "解析目标", "status": "done", "detail": intent_label(context.get("intent"))},
        {"id": "db", "label": "发现企业线索", "status": "done", "detail": db_detail},
        {"id": "policy", "label": "匹配政策机会", "status": "done", "detail": f"{len(context.get('policies') or [])} 条政策，{len(context.get('resources') or [])} 项资源"},
        {"id": "graph", "label": "核验产业位置", "status": "done", "detail": (context.get("industry") or {}).get("industry", "赛道待确认")},
        {"id": "risk", "label": "识别风险因素", "status": "done", "detail": (context.get("risk") or {}).get("risk_level", "风险待确认")},
        {"id": "llm", "label": "形成招商建议", "status": "done", "detail": "整理推荐理由和下一步推进动作"},
    ]


def build_actions(context: dict[str, Any]) -> list[dict[str, Any]]:
    return [
        {"id": "outline", "label": "生成拜访提纲", "type": "material", "material_type": "outline", "hint": "把判断转成会前沟通结构。"},
        {"id": "wechat", "label": "生成微信话术", "type": "material", "material_type": "wechat", "hint": "适合首次或二次触达。"},
        {"id": "briefing", "label": "生成领导汇报摘要", "type": "material", "material_type": "briefing", "hint": "用于五分钟口头汇报。"},
        {"id": "risk", "label": "生成风险复核清单", "type": "material", "material_type": "risk", "hint": "推进前核验争议项。"},
        {"id": "plan", "label": "生成项目推进计划", "type": "material", "material_type": "plan", "hint": "拆成可执行时间线。"},
    ]


def intent_label(intent: str | None) -> str:
    return {
        "data_inventory": "数据盘点",
        "material_generation": "材料生成",
        "mission_discovery": "线索发现",
        "industry_recommendation": "产业方向推荐",
        "industry_analysis": "产业分析",
        "policy_analysis": "政策分析",
        "company_analysis": "企业研判",
    }.get(intent or "", "招商任务")


def compact_context_for_llm(context: dict[str, Any]) -> dict[str, Any]:
    keys = [
        "task",
        "intent",
        "company_name",
        "enterprise",
        "risk",
        "crm_records",
        "policies",
        "resources",
        "industry",
        "external_intelligence",
        "rag_results",
        "sources",
        "requested_industry",
        "requested_count",
        "candidate_enterprises",
        "discovery_mode",
    ]
    compact = {key: context.get(key) for key in keys}
    if compact.get("rag_results"):
        compact["rag_results"] = compact["rag_results"][:4]
    if compact.get("candidate_enterprises"):
        compact["candidate_enterprises"] = compact["candidate_enterprises"][: context.get("requested_count", 5)]
    return compact


def ranked_companies_for_ui(context: dict[str, Any]) -> list[dict[str, Any]]:
    rows = context.get("candidate_enterprises") or []
    return [
        {
            "rank": index + 1,
            "name": row.get("name"),
            "industry": row.get("industry"),
            "sub_industry": row.get("sub_industry"),
            "score": row.get("_score"),
            "score_parts": row.get("_score_parts") or {},
            "reason": row.get("_rank_reason") or row.get("brief") or "",
            "next_step": "进入拜访准备" if index == 0 else "作为备选线索继续核验",
            "employees": row.get("employees"),
            "financing_stage": row.get("financing_stage"),
            "revenue_range": row.get("revenue_range"),
            "patents": row.get("patents"),
        }
        for index, row in enumerate(rows[: context.get("requested_count", 5)])
    ]


def draft_report(context: dict[str, Any], message: str = "") -> dict[str, Any]:
    ranked = ranked_companies_for_ui(context)
    enterprise = context.get("enterprise") or {}
    risk = context.get("risk") or {}
    top_score = ranked[0].get("score") if ranked else enterprise.get("_score", "-")
    summary_target = "、".join([item.get("name", "") for item in ranked[: min(3, len(ranked))]]) if ranked else enterprise.get("name", "")
    return {
        "verdict": "已完成检索，等待最终研判",
        "summary": message or f"已基于本地企业库、政策库、产业图谱与风险数据完成业务底稿，并按任务要求返回 {len(ranked) or 1} 家候选。当前重点线索为：{summary_target}。",
        "confidence": "待模型确认",
        "metrics": {
            "match_score": top_score or "-",
            "risk_level": risk.get("risk_level", "-"),
            "recommended_action": "查看推荐依据，继续生成材料或补充筛选约束",
        },
        "sections": [
            {
                "id": "shortlist",
                "title": "候选线索",
                "body": "以下名单来自本地企业库和任务约束匹配，不是前端写死结果。",
                "bullets": [f"{item.get('rank')}. {item.get('name')}：{item.get('reason')}" for item in ranked],
            },
            {
                "id": "evidence",
                "title": "已命中依据",
                "body": "系统已完成结构化数据读取，并保留每条结论相关的来源卡片。",
                "bullets": [item.get("title", "") for item in context.get("evidence", [])[:6]],
            },
        ],
        "policy_matches": [policy.get("name", "") for policy in (context.get("policies") or [])[:5]],
        "action_plan": ["生成拜访提纲", "生成微信话术", "生成领导汇报摘要", "继续补充筛选条件"],
        "ranked_companies": ranked,
        "sources_used": [item["id"] for item in context.get("evidence", [])],
        "draft": True,
    }


def call_llm(messages: list[dict[str, str]], max_tokens: int = 3600) -> str:
    api_key = get_api_key()
    if not api_key:
        raise RuntimeError("未配置 DASHSCOPE_API_KEY，且 original V0/api_key.txt 不可用")
    payload = {
        "model": LLM_MODEL,
        "messages": messages,
        "temperature": 0.28,
        "max_tokens": max_tokens,
    }
    req = urllib.request.Request(
        f"{LLM_API_BASE.rstrip('/')}/chat/completions",
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=LLM_TIMEOUT) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"LLM HTTP {exc.code}: {body[:500]}") from exc
    except Exception as exc:
        raise RuntimeError(f"LLM 调用失败：{exc}") from exc

    return data["choices"][0]["message"]["content"].strip()


def call_llm_stream(messages: list[dict[str, str]], on_chunk, max_tokens: int = 3600) -> str:
    api_key = get_api_key()
    if not api_key:
        raise RuntimeError("未配置 DASHSCOPE_API_KEY，且 original V0/api_key.txt 不可用")
    payload = {
        "model": LLM_MODEL,
        "messages": messages,
        "temperature": 0.28,
        "max_tokens": max_tokens,
        "stream": True,
        "stream_options": {"include_usage": True},
    }
    req = urllib.request.Request(
        f"{LLM_API_BASE.rstrip('/')}/chat/completions",
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    chunks: list[str] = []
    try:
        with urllib.request.urlopen(req, timeout=LLM_TIMEOUT) as resp:
            for line_bytes in resp:
                line = line_bytes.decode("utf-8", errors="replace").strip()
                if not line or not line.startswith("data:"):
                    continue
                data_str = line[len("data:"):].strip()
                if data_str == "[DONE]":
                    break
                try:
                    event = json.loads(data_str)
                except json.JSONDecodeError:
                    continue
                delta = (event.get("choices") or [{}])[0].get("delta") or {}
                content = delta.get("content")
                if content:
                    chunks.append(content)
                    on_chunk(content)
    except urllib.error.HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"LLM HTTP {exc.code}: {body[:500]}") from exc
    except Exception as exc:
        raise RuntimeError(f"LLM 流式调用失败：{exc}") from exc

    return "".join(chunks).strip()


def parse_json_object(text: str) -> dict[str, Any] | None:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?", "", cleaned).strip()
        cleaned = re.sub(r"```$", "", cleaned).strip()
    try:
        return json.loads(cleaned)
    except Exception:
        pass
    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start >= 0 and end > start:
        try:
            return json.loads(cleaned[start : end + 1])
        except Exception:
            return None
    return None


def generate_report(context: dict[str, Any]) -> dict[str, Any]:
    system = (
        "你是资深产业园区招商顾问。必须只基于提供的本地数据、RAG片段和风险信息生成招商研判。"
        "不要编造不存在的数据。输出必须是 JSON，不要 Markdown。"
    )
    user = {
        "instruction": (
            "生成一份可用于招商工作台展示的结构化研判。"
            "结论要先行，语言克制专业。"
            "如果 intent 是 industry_recommendation，必须说明当前企业来自候选企业排序，不得偏离 requested_industry。"
            "如果 intent 是 mission_discovery，必须像招商 Agent 一样先给出候选名单、排序逻辑和下一步核验动作，不要要求用户先提供企业名。"
            "如果 local_context.requested_count 存在，ranked_companies 必须尽量返回对应数量，不要擅自改成 5 家。"
            "如果证据不足，要明确写待核验事项。"
        ),
        "required_schema": {
            "verdict": "值得重点推进/谨慎推进/暂不建议推进",
            "summary": "120字以内核心结论",
            "confidence": "高/中/低",
            "metrics": {
                "match_score": "0-100整数",
                "risk_level": "低风险/中等风险/高风险",
                "recommended_action": "下一步动作",
            },
            "sections": [
                {"id": "profile", "title": "企业画像", "body": "段落", "bullets": ["要点"]},
                {"id": "industry", "title": "产业链位置", "body": "段落", "bullets": ["要点"]},
                {"id": "policy", "title": "政策抓手", "body": "段落", "bullets": ["要点"]},
                {"id": "risk", "title": "风险与待核验", "body": "段落", "bullets": ["要点"]},
                {"id": "action", "title": "推进建议", "body": "段落", "bullets": ["要点"]},
            ],
            "policy_matches": ["政策匹配要点"],
            "action_plan": ["可执行动作"],
            "ranked_companies": [
                {
                    "rank": "序号",
                    "name": "企业名",
                    "score": "综合评分",
                    "reason": "推荐原因",
                    "next_step": "下一步核验或推进动作",
                }
            ],
            "sources_used": ["引用的 evidence id"],
        },
        "local_context": compact_context_for_llm(context),
        "evidence": context.get("evidence", []),
    }
    text = call_llm(
        [
            {"role": "system", "content": system},
            {"role": "user", "content": json.dumps(user, ensure_ascii=False)},
        ]
    )
    parsed = parse_json_object(text)
    if not parsed:
        return {
            "verdict": "模型返回未结构化结果",
            "summary": "模型返回了真实文本，但未按 JSON 结构输出。以下保留原文用于排查。",
            "confidence": "低",
            "metrics": {
                "match_score": "-",
                "risk_level": context.get("risk", {}).get("risk_level", "-"),
                "recommended_action": "重新生成",
            },
            "sections": [{"id": "raw", "title": "模型原文", "body": text, "bullets": []}],
            "policy_matches": [],
            "action_plan": [],
            "ranked_companies": ranked_companies_for_ui(context),
            "sources_used": [item["id"] for item in context.get("evidence", [])],
        }
    if context.get("candidate_enterprises"):
        parsed["ranked_companies"] = ranked_companies_for_ui(context)
    return parsed


def generate_report_stream(context: dict[str, Any], on_chunk) -> dict[str, Any]:
    system = (
        "你是资深产业园区招商顾问。必须只基于提供的本地数据、RAG片段和风险信息生成招商研判。"
        "不要编造不存在的数据。输出必须是 JSON，不要 Markdown。"
    )
    user = {
        "instruction": (
            "生成一份可用于招商工作台展示的结构化研判。"
            "结论要先行，语言克制专业。"
            "如果 intent 是 industry_recommendation，必须说明当前企业来自候选企业排序，不得偏离 requested_industry。"
            "如果 intent 是 mission_discovery，必须像招商 Agent 一样先给出候选名单、排序逻辑和下一步核验动作，不要要求用户先提供企业名。"
            "如果 local_context.requested_count 存在，ranked_companies 必须尽量返回对应数量，不要擅自改成 5 家。"
            "如果证据不足，要明确写待核验事项。"
        ),
        "required_schema": {
            "verdict": "值得重点推进/谨慎推进/暂不建议推进",
            "summary": "120字以内核心结论",
            "confidence": "高/中/低",
            "metrics": {
                "match_score": "0-100整数",
                "risk_level": "低风险/中等风险/高风险",
                "recommended_action": "下一步动作",
            },
            "sections": [
                {"id": "profile", "title": "企业画像", "body": "段落", "bullets": ["要点"]},
                {"id": "industry", "title": "产业链位置", "body": "段落", "bullets": ["要点"]},
                {"id": "policy", "title": "政策抓手", "body": "段落", "bullets": ["要点"]},
                {"id": "risk", "title": "风险与待核验", "body": "段落", "bullets": ["要点"]},
                {"id": "action", "title": "推进建议", "body": "段落", "bullets": ["要点"]},
            ],
            "policy_matches": ["政策匹配要点"],
            "action_plan": ["可执行动作"],
            "ranked_companies": [
                {
                    "rank": "序号",
                    "name": "企业名",
                    "score": "综合评分",
                    "reason": "推荐原因",
                    "next_step": "下一步核验或推进动作",
                }
            ],
            "sources_used": ["引用的 evidence id"],
        },
        "local_context": compact_context_for_llm(context),
        "evidence": context.get("evidence", []),
    }
    text = call_llm_stream(
        [
            {"role": "system", "content": system},
            {"role": "user", "content": json.dumps(user, ensure_ascii=False)},
        ],
        on_chunk=on_chunk,
    )
    parsed = parse_json_object(text)
    if not parsed:
        return {
            "verdict": "模型返回未结构化结果",
            "summary": "模型返回了真实文本，但未按 JSON 结构输出。以下保留原文用于排查。",
            "confidence": "低",
            "metrics": {
                "match_score": "-",
                "risk_level": context.get("risk", {}).get("risk_level", "-"),
                "recommended_action": "重新生成",
            },
            "sections": [{"id": "raw", "title": "模型原文", "body": text, "bullets": []}],
            "policy_matches": [],
            "action_plan": [],
            "ranked_companies": ranked_companies_for_ui(context),
            "sources_used": [item["id"] for item in context.get("evidence", [])],
        }
    if context.get("candidate_enterprises"):
        parsed["ranked_companies"] = ranked_companies_for_ui(context)
    return parsed


def generate_material(material_type: str, report: dict[str, Any], context: dict[str, Any]) -> dict[str, Any]:
    label = MATERIAL_LABELS.get(material_type, material_type or "招商材料")
    system = (
        "你是资深招商材料写作顾问。必须基于报告、证据和本地数据生成可直接使用的招商材料。"
        "不得编造新事实。输出 JSON，不要 Markdown。"
    )
    user = {
        "instruction": f"生成《{label}》。语气专业、克制、可直接用于招商工作。",
        "required_schema": {
            "title": label,
            "audience": "对象",
            "content": ["分段内容"],
            "source_notes": ["依据"],
        },
        "report": report,
        "local_context": compact_context_for_llm(context),
        "evidence": context.get("evidence", []),
    }
    text = call_llm(
        [
            {"role": "system", "content": system},
            {"role": "user", "content": json.dumps(user, ensure_ascii=False)},
        ],
        max_tokens=2400,
    )
    parsed = parse_json_object(text)
    if not parsed:
        return {"title": label, "audience": "未解析", "content": [text], "source_notes": ["真实 LLM 输出，未结构化解析"]}
    return parsed


def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value)


def _setup_doc_style(doc: Any, normal_font: str = "仿宋", normal_size: int = 16) -> Any:
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn

    for section in doc.sections:
        section.top_margin = Cm(2.8)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = normal_font
    style.font.size = Pt(normal_size)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), normal_font)
    pf = style.paragraph_format
    pf.line_spacing = Pt(28)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(normal_size * 2)
    return style


def _add_doc_title(doc: Any, text: str) -> Any:
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = Pt(36)
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = "宋体"
    run.font.size = Pt(22)
    run.bold = True
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p


def _add_doc_heading(doc: Any, text: str, font_name: str = "黑体", size: int = 16) -> Any:
    from docx.shared import Pt
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = Pt(28)
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = True
    run.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    return p


def _add_doc_subheading(doc: Any, text: str) -> Any:
    return _add_doc_heading(doc, text, font_name="楷体", size=15)


def _add_doc_body(doc: Any, text: str, indent: bool = True) -> Any:
    from docx.shared import Pt
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = Pt(28)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(32) if indent else Pt(0)
    run = p.add_run(text)
    run.font.name = "仿宋"
    run.font.size = Pt(16)
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    return p


def _add_doc_bullet(doc: Any, text: str) -> Any:
    from docx.shared import Pt
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = Pt(28)
    pf.left_indent = Pt(32)
    pf.first_line_indent = Pt(-16)
    run = p.add_run(f"— {text}")
    run.font.name = "仿宋"
    run.font.size = Pt(16)
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    return p


def _add_doc_footer(doc: Any) -> None:
    from docx.shared import Pt
    from docx.oxml.ns import qn

    _add_doc_body(doc, "", indent=False)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = Pt(28)
    pf.space_before = Pt(16)
    pf.first_line_indent = Pt(0)
    pf.alignment = 1
    run = p.add_run("由 ParkFlow AI 生成，仅供参考。")
    run.font.name = "仿宋"
    run.font.size = Pt(10)
    run.italic = True
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")


def _clean_text(value: Any) -> str:
    text = _safe_text(value)
    text = text.replace("**", "").replace("__", "")
    text = text.replace("```", "").replace("`", "")
    text = text.replace("——", "——")
    return text.strip()


def _chinese_num(n: int) -> str:
    nums = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    if n <= 10:
        return nums[n]
    if n < 20:
        return f"十{nums[n - 10]}"
    return str(n)


def _build_report_table(doc: Any, ranked: list[dict[str, Any]]) -> None:
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not ranked:
        return
    table = doc.add_table(rows=1 + min(len(ranked), 10), cols=5)
    table.style = "Table Grid"
    headers = ["序号", "企业名称", "产业方向", "匹配分", "推荐原因"]
    widths = [Cm(1.0), Cm(3.5), Cm(2.5), Cm(1.5), Cm(7.5)]
    for i, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = width
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = Pt(22)
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(header)
        run.font.name = "黑体"
        run.font.size = Pt(10)
        run.bold = True
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    for idx, company in enumerate(ranked[:10]):
        if not isinstance(company, dict):
            continue
        row = table.rows[idx + 1]
        values = [
            str(idx + 1),
            _clean_text(company.get("name", "")),
            _clean_text(company.get("industry", company.get("sub_industry", ""))),
            str(company.get("score", "-")),
            _clean_text(company.get("reason", ""))[:60],
        ]
        for col, value in enumerate(values):
            cell = row.cells[col]
            cell.width = widths[col]
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = Pt(20)
            p.paragraph_format.first_line_indent = Pt(0)
            run = p.add_run(value)
            run.font.name = "仿宋"
            run.font.size = Pt(9)
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    _add_doc_body(doc, "", indent=False)


def build_report_docx(report: dict[str, Any], context: dict[str, Any]) -> bytes:
    from docx import Document

    doc = Document()
    _setup_doc_style(doc)
    _add_doc_title(doc, "园区重点招商企业推荐建议")

    verdict = _clean_text(report.get("verdict", ""))
    summary = _clean_text(report.get("summary", ""))
    metrics = report.get("metrics") or {}
    sections = report.get("sections") or []
    ranked = report.get("ranked_companies") or context.get("candidate_enterprises") or []

    _add_doc_heading(doc, "一、总体判断")
    _add_doc_body(doc, f"经综合研判，{verdict}。{summary}")

    _add_doc_heading(doc, "二、筛选逻辑")
    candidate_count = len(ranked)
    industries = list({_safe_text(c.get("industry", "")) for c in ranked[:5] if isinstance(c, dict)})
    logic_text = f"基于本地企业库、产业图谱、政策匹配与风险数据，从候选池中筛选出 {candidate_count} 家企业。重点考量产业方向匹配度（{', '.join(industries[:3]) or '多元方向'}）、企业成长阶段、租金承载力、技术积累与风险安全等维度，综合排序形成推荐名单。"
    _add_doc_body(doc, logic_text)

    _add_doc_heading(doc, "三、推荐企业清单")
    _build_report_table(doc, ranked)

    section_counter = 4
    profile_section = next((s for s in sections if isinstance(s, dict) and s.get("id") == "profile"), None)
    if profile_section:
        _add_doc_heading(doc, f"{_chinese_num(section_counter)}、{_clean_text(profile_section.get('title', '重点企业说明'))}")
        section_counter += 1
        body = _clean_text(profile_section.get("body", ""))
        if body:
            _add_doc_body(doc, body)
        for bullet in (profile_section.get("bullets") or []):
            _add_doc_bullet(doc, _clean_text(bullet))
    elif ranked:
        _add_doc_heading(doc, f"{_chinese_num(section_counter)}、重点企业说明")
        section_counter += 1
        top = ranked[0] if isinstance(ranked[0], dict) else {}
        _add_doc_body(doc, f"首推企业为{_clean_text(top.get('name', ''))}，"
                       f"该企业聚焦{_clean_text(top.get('industry', ''))}赛道，"
                       f"匹配分{_safe_text(top.get('score', '-'))}分，"
                       f"具备较强的产业契合度与落地潜力。")

    risk_section = next((s for s in sections if isinstance(s, dict) and s.get("id") == "risk"), None)
    if risk_section:
        _add_doc_heading(doc, f"{_chinese_num(section_counter)}、风险与待核验事项")
        section_counter += 1
        body = _clean_text(risk_section.get("body", ""))
        if body:
            _add_doc_body(doc, body)
        for bullet in (risk_section.get("bullets") or []):
            _add_doc_bullet(doc, _clean_text(bullet))
    else:
        _add_doc_heading(doc, f"{_chinese_num(section_counter)}、风险与待核验事项")
        section_counter += 1
        risk_level = _safe_text(metrics.get("risk_level", "待确认"))
        _add_doc_body(doc, f"当前风险等级为{risk_level}。建议在正式推进前，逐项核验企业资质、经营状态与落地意愿。")

    action_section = next((s for s in sections if isinstance(s, dict) and s.get("id") == "action"), None)
    action_plan = report.get("action_plan") or []
    _add_doc_heading(doc, f"{_chinese_num(section_counter)}、下一步推进建议")
    section_counter += 1
    if action_section:
        body = _clean_text(action_section.get("body", ""))
        if body:
            _add_doc_body(doc, body)
    if action_plan:
        for idx, action in enumerate(action_plan, 1):
            _add_doc_body(doc, f"（{_chinese_num(idx)}）{_clean_text(action)}")
    if not action_plan and not action_section:
        _add_doc_body(doc, "建议优先推进排名靠前的企业，同步开展资质核验与政策匹配，择机启动拜访邀约。")

    policy_matches = report.get("policy_matches") or []
    if policy_matches:
        _add_doc_heading(doc, f"{_chinese_num(section_counter)}、可匹配政策资源")
        section_counter += 1
        for policy in policy_matches:
            _add_doc_bullet(doc, _clean_text(policy))

    evidence = context.get("evidence") or []
    if evidence:
        _add_doc_heading(doc, "附件：判断依据")
        for item in evidence[:8]:
            if not isinstance(item, dict):
                continue
            src = f"{_clean_text(item.get('title'))}（{_safe_text(item.get('source'))}）"
            _add_doc_bullet(doc, src)

    _add_doc_footer(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_material_docx(material: dict[str, Any], material_type: str) -> bytes:
    from docx import Document

    doc = Document()
    _setup_doc_style(doc)

    title = _clean_text(material.get("title")) or MATERIAL_LABELS.get(material_type, "招商材料")
    audience = _clean_text(material.get("audience"))
    content = material.get("content") or []
    source_notes = material.get("source_notes") or []

    if material_type == "outline":
        company_name = _clean_text(material.get("company_name") or audience or "企业")
        doc_title = f"关于赴{company_name}开展招商拜访的提纲"
        _add_doc_title(doc, doc_title)
        _add_doc_heading(doc, "一、拜访目的")
        _add_doc_body(doc, f"了解{company_name}的发展现状、核心产品与市场布局，推介光谷智创园的产业生态、政策优势与空间资源，探讨双方在技术研发、产业协同、项目落地等领域的合作机会。")
        _add_doc_heading(doc, "二、企业背景预研")
        _add_doc_body(doc, f"提前研读{company_name}的公开信息：主营业务、融资阶段、团队背景、专利布局、近期动态。梳理其与园区现有产业集群的潜在协同点，准备针对性问题清单。")
        _add_doc_heading(doc, "三、园区推介要点")
        _add_doc_body(doc, "重点介绍园区产业集聚效应（45家高新技术企业入驻）、科研支撑（国家级实验室与中试基地）、人才保障（周边3所985高校）、交通与生活配套等核心优势，结合企业实际情况侧重呈现相关政策资源。")
        _add_doc_heading(doc, "四、政策对接准备")
        _add_doc_body(doc, "根据企业资质与需求，提前准备可匹配的政策清单与申报条件，编制《政策兑现路径图》，确保拜访中能够清晰传达园区的实质性支持。")
        _add_doc_heading(doc, "五、沟通提纲")
        for idx, para in enumerate(content[:8], 1):
            _add_doc_body(doc, f"（{_chinese_num(idx)}）{_clean_text(para)}")
        if not content:
            _add_doc_body(doc, "（一）了解企业当前发展阶段、核心诉求与布局规划。")
            _add_doc_body(doc, "（二）介绍园区产业生态、空间资源与政策配套。")
            _add_doc_body(doc, "（三）探讨潜在合作模式与落地可行性。")
        _add_doc_heading(doc, "六、需重点确认事项")
        _add_doc_bullet(doc, "企业是否已获得高新技术企业等相关资质认定")
        _add_doc_bullet(doc, "核心团队跨区域搬迁或设立分支机构的意愿与条件")
        _add_doc_bullet(doc, "企业当前对办公/厂房空间的面积与配套需求")
        _add_doc_heading(doc, "七、后续跟进节点")
        _add_doc_body(doc, "拜访结束后24小时内发送会议纪要与政策匹配清单；48小时内确认双方下一步对接人及时间表；一周内完成内部研判并形成推进方案。")
        _add_doc_heading(doc, "八、附件清单")
        _add_doc_bullet(doc, "园区宣传册（电子版）")
        _add_doc_bullet(doc, "政策匹配清单与企业资质对照表")
        _add_doc_bullet(doc, "园区空间资源一览（含户型图与实景照片）")
        if source_notes:
            for note in source_notes[:3]:
                _add_doc_bullet(doc, _clean_text(note))

    elif material_type == "wechat":
        company_name = _clean_text(material.get("company_name") or audience or "企业")
        doc_title = f"{company_name}招商跟进话术"
        _add_doc_title(doc, doc_title)
        _add_doc_heading(doc, "一、首次触达话术")
        _add_doc_body(doc, f"{company_name}您好，我是光谷智创园招商负责人。关注到贵司在{_clean_text(material.get('industry', '相关领域'))}的突出表现，我园已集聚45家高新技术企业，配备国家级实验室与中试基地，可为企业提供从空间到政策的全方位支持。诚邀您择机来园实地考察，期待深入交流。")
        _add_doc_heading(doc, "二、二次跟进话术")
        _add_doc_body(doc, f"上次与您沟通后，我们针对{company_name}的发展方向梳理了匹配的政策清单与空间方案。园区在产业集聚、科研支撑与人才保障方面具有明显优势，期待能与贵司进一步探讨合作可能。")
        _add_doc_heading(doc, "三、政策引导话术")
        _add_doc_body(doc, "根据贵司的资质条件，可叠加享受高新技术企业入驻奖励（3年免租）、专精特新企业奖励（50万元）以及人才引进安家补贴（博士30万/硕士10万）。我们可协助您完成全流程申报。")
        _add_doc_heading(doc, "四、异议应对")
        _add_doc_body(doc, "若企业表达对园区区位、配套或成本的疑虑：理解您的考量。光谷智创园毗邻华中科技大学与武汉大学，人才资源丰富；交通方面，地铁、高速、高铁三位一体；空间方面，从50平方米孵化工位到5000平方米大型厂房均可灵活配置。建议实地考察后再做评估。")
        _add_doc_heading(doc, "五、促成邀约")
        _add_doc_body(doc, f"真诚邀请{company_name}团队来园实地考察，我们将安排专人全程陪同，参观园区核心设施、对接已入驻龙头企业，并就具体政策支持进行一对一沟通。如方便，请告知贵司近期可安排的时间。")
        _add_doc_heading(doc, "六、注意事项")
        _add_doc_bullet(doc, "话术使用时应根据实际沟通场景灵活调整，保持专业、克制的语调")
        _add_doc_bullet(doc, "首次触达建议通过微信或电话进行，避免群发式消息")
        _add_doc_bullet(doc, "跟进节奏控制在每周1-2次，避免过度打扰")
        if source_notes:
            for note in source_notes[:3]:
                _add_doc_bullet(doc, _clean_text(note))

    elif material_type == "briefing":
        doc_title = "招商工作汇报摘要"
        _add_doc_title(doc, doc_title)
        _add_doc_heading(doc, "一、本期重点线索")
        if content:
            for para in content[:3]:
                _add_doc_body(doc, _clean_text(para))
        else:
            _add_doc_body(doc, f"根据招商任务要求，已完成企业筛选与研判，识别出具有较高落地潜力的候选企业，现将研判结果与推进建议汇报如下。")
        _add_doc_heading(doc, "二、候选企业概况")
        _add_doc_body(doc, _clean_text(material.get("summary", "详见招商研判报告中的推荐企业清单与排序依据。")))
        _add_doc_heading(doc, "三、政策与资源匹配")
        if source_notes:
            for note in source_notes[:3]:
                _add_doc_bullet(doc, _clean_text(note))
        _add_doc_heading(doc, "四、下一步工作安排")
        _add_doc_body(doc, "一是启动首推企业的定向拜访邀约；二是推进政策匹配与空间资源预配；三是持续跟踪候选企业动态，滚动更新推荐名单。")
        _add_doc_heading(doc, "五、需要协调事项")
        _add_doc_body(doc, "请领导协调相关部门，为本次招商推进提供必要的政策解读、空间调配与接待支持。")
        if audience:
            _add_doc_heading(doc, "六、汇报对象")
            _add_doc_body(doc, audience)

    else:
        _add_doc_title(doc, title)
        if audience:
            from docx.shared import Pt
            from docx.oxml.ns import qn
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = Pt(28)
            p.paragraph_format.first_line_indent = Pt(0)
            run = p.add_run(f"对象：{audience}")
            run.font.name = "楷体"
            run.font.size = Pt(14)
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
        if isinstance(content, list):
            for para in content:
                _add_doc_body(doc, _clean_text(para))
        else:
            _add_doc_body(doc, _clean_text(content))
        if source_notes:
            _add_doc_heading(doc, "依据")
            for note in source_notes:
                _add_doc_bullet(doc, _clean_text(note))

    _add_doc_footer(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_PDF_FONT_NAME = "msyh"
_PDF_FONT_PATH = "C:/Windows/Fonts/msyh.ttc"
_PDF_FONT_BOLD = "C:/Windows/Fonts/msyhbd.ttc"


def _build_pdf_styles() -> Any:
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.units import mm
    from reportlab.lib import colors

    return {
        "title": ParagraphStyle("pdf_title", fontName=_PDF_FONT_NAME, fontSize=18, leading=24, alignment=TA_CENTER, spaceAfter=8 * mm),
        "subtitle": ParagraphStyle("pdf_sub", fontName=_PDF_FONT_NAME, fontSize=13, leading=18, spaceAfter=4 * mm),
        "heading": ParagraphStyle("pdf_h", fontName=_PDF_FONT_NAME, fontSize=13, leading=18, spaceBefore=6 * mm, spaceAfter=3 * mm),
        "body": ParagraphStyle("pdf_body", fontName=_PDF_FONT_NAME, fontSize=10, leading=16, alignment=TA_LEFT, spaceAfter=2 * mm),
        "small": ParagraphStyle("pdf_small", fontName=_PDF_FONT_NAME, fontSize=8, leading=12, alignment=TA_CENTER),
        "table_header": ParagraphStyle("pdf_th", fontName=_PDF_FONT_NAME, fontSize=9, leading=13, alignment=TA_CENTER),
        "table_cell": ParagraphStyle("pdf_td", fontName=_PDF_FONT_NAME, fontSize=9, leading=13),
        "footer": ParagraphStyle("pdf_foot", fontName=_PDF_FONT_NAME, fontSize=8, leading=12, alignment=TA_CENTER, textColor=colors.gray),
    }


def _register_pdf_font() -> None:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    try:
        pdfmetrics.registerFont(TTFont(_PDF_FONT_NAME, _PDF_FONT_PATH))
        pdfmetrics.registerFont(TTFont(_PDF_FONT_NAME + "-Bold", _PDF_FONT_BOLD))
    except Exception:
        pass


def build_report_pdf(report: dict[str, Any], context: dict[str, Any]) -> bytes:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.units import mm
    from reportlab.lib import colors

    _register_pdf_font()
    S = _build_pdf_styles()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=(210 * mm, 297 * mm),
                           leftMargin=25 * mm, rightMargin=20 * mm,
                           topMargin=25 * mm, bottomMargin=20 * mm)
    story = []

    story.append(Paragraph("企业招商研判报告", S["title"]))
    story.append(Spacer(1, 4 * mm))

    verdict = _safe_text(report.get("verdict", ""))
    summary = _safe_text(report.get("summary", ""))
    metrics = report.get("metrics") or {}

    core_style = ParagraphStyle("core", parent=S["subtitle"], fontName=_PDF_FONT_NAME + "-Bold")
    story.append(Paragraph(f"核心判断：{verdict}", core_style))
    if summary:
        story.append(Paragraph(summary, S["body"]))

    story.append(Paragraph("关键指标", S["heading"]))
    story.append(Paragraph(
        f"匹配度：{_safe_text(metrics.get('match_score', '-'))}　　"
        f"风险等级：{_safe_text(metrics.get('risk_level', '-'))}　　"
        f"建议动作：{_safe_text(metrics.get('recommended_action', '-'))}",
        S["body"]
    ))

    sections = report.get("sections") or []
    for section in sections:
        if not isinstance(section, dict):
            continue
        sec_title = _safe_text(section.get("title"))
        if not sec_title:
            continue
        h_style = ParagraphStyle("sec_h", parent=S["heading"], fontName=_PDF_FONT_NAME + "-Bold")
        story.append(Paragraph(sec_title, h_style))
        body = _safe_text(section.get("body"))
        if body:
            story.append(Paragraph(body, S["body"]))
        bullets = section.get("bullets") or []
        for bullet in bullets:
            story.append(Paragraph(f"• {_safe_text(bullet)}", S["body"]))

    ranked = report.get("ranked_companies") or context.get("candidate_enterprises") or []
    if ranked:
        story.append(Paragraph("候选企业排名", S["heading"]))
        headers = ["#", "企业名称", "匹配分", "推荐理由", "下一步建议"]
        header_row = [Paragraph(h, S["table_header"]) for h in headers]
        table_data = [header_row]
        for idx, company in enumerate(ranked[:8]):
            if not isinstance(company, dict):
                continue
            row = [
                Paragraph(str(idx + 1), S["table_header"]),
                Paragraph(_safe_text(company.get("name"))[:18], S["table_cell"]),
                Paragraph(str(company.get("score", "-")), S["table_header"]),
                Paragraph(_safe_text(company.get("reason", ""))[:50], S["table_cell"]),
                Paragraph(_safe_text(company.get("next_step", ""))[:24], S["table_cell"]),
            ]
            table_data.append(row)
        col_w = [12 * mm, 42 * mm, 16 * mm, 70 * mm, 36 * mm]
        tbl = Table(table_data, colWidths=col_w, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.append(tbl)

    policy_matches = report.get("policy_matches") or []
    if policy_matches:
        story.append(Paragraph("政策匹配", S["heading"]))
        for policy in policy_matches:
            story.append(Paragraph(f"• {_safe_text(policy)}", S["body"]))

    action_plan = report.get("action_plan") or []
    if action_plan:
        story.append(Paragraph("推进计划", S["heading"]))
        for idx, action in enumerate(action_plan, 1):
            story.append(Paragraph(f"{idx}. {_safe_text(action)}", S["body"]))

    evidence = context.get("evidence") or []
    if evidence:
        story.append(Paragraph("判断依据（资料引用）", S["heading"]))
        for item in evidence:
            if not isinstance(item, dict):
                continue
            src = f"{_safe_text(item.get('title'))} ({_safe_text(item.get('source'))})"
            story.append(Paragraph(f"• {src}", S["small"]))

    story.append(Spacer(1, 6 * mm))
    story.append(Paragraph("由 ParkFlow AI 生成，仅供参考，不作为唯一决策依据。", S["footer"]))

    doc.build(story)
    return buf.getvalue()


def build_material_pdf(material: dict[str, Any], material_type: str) -> bytes:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.units import mm

    _register_pdf_font()
    S = _build_pdf_styles()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=(210 * mm, 297 * mm),
                           leftMargin=25 * mm, rightMargin=20 * mm,
                           topMargin=25 * mm, bottomMargin=20 * mm)
    story = []

    label = MATERIAL_LABELS.get(material_type, material_type or "招商材料")
    title = _safe_text(material.get("title")) or label
    story.append(Paragraph(title, S["title"]))
    story.append(Spacer(1, 6 * mm))

    audience = _safe_text(material.get("audience"))
    if audience:
        aud_style = ParagraphStyle("aud", parent=S["body"], fontName=_PDF_FONT_NAME + "-Bold")
        story.append(Paragraph(f"对象：{audience}", aud_style))
        story.append(Spacer(1, 3 * mm))

    content = material.get("content") or []
    if isinstance(content, list):
        for para in content:
            story.append(Paragraph(_safe_text(para), S["body"]))
            story.append(Spacer(1, 1.5 * mm))
    else:
        story.append(Paragraph(_safe_text(content), S["body"]))

    source_notes = material.get("source_notes") or []
    if source_notes:
        story.append(Spacer(1, 3 * mm))
        story.append(Paragraph("依据", S["heading"]))
        for note in source_notes:
            story.append(Paragraph(f"• {_safe_text(note)}", S["small"]))

    story.append(Spacer(1, 6 * mm))
    story.append(Paragraph("由 ParkFlow AI 生成，可直接用于招商工作。", S["footer"]))

    doc.build(story)
    return buf.getvalue()


class Handler(BaseHTTPRequestHandler):
    server_version = "ParkFlowMVP/0.2"

    def log_message(self, fmt: str, *args: Any) -> None:
        sys.stderr.write("%s - - [%s] %s\n" % (self.client_address[0], self.log_date_time_string(), fmt % args))

    def send_cors_headers(self) -> None:
        origin = (self.headers.get("Origin") or "").rstrip("/")
        if not ALLOWED_ORIGINS:
            return
        if "*" in ALLOWED_ORIGINS:
            self.send_header("Access-Control-Allow-Origin", "*")
        elif origin in ALLOWED_ORIGINS:
            self.send_header("Access-Control-Allow-Origin", origin)
            self.send_header("Vary", "Origin")
        else:
            return
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type, Authorization")
        self.send_header("Access-Control-Max-Age", "86400")

    def do_OPTIONS(self) -> None:
        self.send_response(204)
        self.send_cors_headers()
        self.send_header("Content-Length", "0")
        self.end_headers()

    def send_json(self, payload: Any, status: int = 200) -> None:
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Cache-Control", "no-store")
        self.send_cors_headers()
        self.end_headers()
        self.wfile.write(data)

    def send_stream_event(self, event: dict[str, Any]) -> None:
        data = (json.dumps(event, ensure_ascii=False) + "\n").encode("utf-8")
        self.wfile.write(data)
        self.wfile.flush()

    def start_ndjson_stream(self) -> None:
        self.send_response(200)
        self.send_header("Content-Type", "application/x-ndjson; charset=utf-8")
        self.send_header("Cache-Control", "no-store")
        self.send_cors_headers()
        self.end_headers()

    def read_body(self) -> dict[str, Any]:
        length = int(self.headers.get("Content-Length") or 0)
        raw = self.rfile.read(length).decode("utf-8") if length else "{}"
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            return {}

    def do_GET(self) -> None:
        if self.path.startswith("/api/health"):
            self.send_json(
                {
                    "ok": True,
                    "db_exists": DB_PATH.exists(),
                    "original_v0_readonly": True,
                    "llm_configured": bool(get_api_key()),
                    "model": LLM_MODEL,
                    "capabilities": [
                        "local_enterprise_db",
                        "crm_records",
                        "policy_db",
                        "industry_graph",
                        "risk_data",
                        "rag_best_effort",
                    ],
                }
            )
            return
        if self.path.startswith("/api/stats"):
            self.send_json({"ok": True, "stats": data_stats()})
            return
        if self.path.startswith("/api/companies"):
            query = ""
            if "?" in self.path:
                query = urllib.parse.parse_qs(urllib.parse.urlparse(self.path).query).get("q", [""])[0]
            if query:
                rows = db_rows(
                    "SELECT id,name,industry,sub_industry,employees,financing_stage,region FROM enterprises "
                    "WHERE name LIKE ? OR industry LIKE ? OR sub_industry LIKE ? LIMIT 20",
                    (f"%{query}%", f"%{query}%", f"%{query}%"),
                )
            else:
                rows = db_rows(
                    "SELECT id,name,industry,sub_industry,employees,financing_stage,region FROM enterprises "
                    "ORDER BY id LIMIT 20"
                )
            self.send_json({"ok": True, "companies": rows})
            return
        self.serve_static()

    def handle_original_agent_stream(self, task: str, started: float) -> None:
        self.start_ndjson_stream()
        self.send_stream_event(
            {
                "event": "agent_state",
                "mode": "conversation",
                "status": "active",
                "label": "正在回复",
                "detail": "正在理解你的问题",
                "elapsed_ms": 0,
            }
        )
        chunks: list[str] = []

        def on_chunk(piece: str) -> None:
            chunks.append(piece)
            self.send_stream_event(
                {
                    "event": "chunk",
                    "content": piece,
                    "elapsed_ms": int((time.time() - started) * 1000),
                }
            )

        def emit_text(text: str) -> None:
            for index in range(0, len(text), 48):
                on_chunk(text[index : index + 48])
                time.sleep(0.02)

        try:
            with _ORIGINAL_AGENT_LOCK:
                old_cwd = os.getcwd()
                try:
                    os.chdir(ORIGINAL_DIR)
                    agent = get_original_agent_unlocked()
                    answer = agent.chat(task, on_stream_chunk=on_chunk)
                finally:
                    os.chdir(old_cwd)
            content = "".join(chunks).strip() or str(answer or "").strip()
            self.send_stream_event(
                {
                    "event": "message",
                    "ok": True,
                    "mode": "conversation",
                    "content": content,
                    "elapsed_ms": int((time.time() - started) * 1000),
                }
            )
        except Exception as exc:
            try:
                answer = call_original_style_chat(task)
                emit_text(answer)
                self.send_stream_event(
                    {
                        "event": "message",
                        "ok": True,
                        "mode": "conversation",
                        "content": answer,
                        "adapter_note": f"original_agent_runtime_unavailable: {exc}",
                        "elapsed_ms": int((time.time() - started) * 1000),
                    }
                )
            except Exception as fallback_exc:
                self.send_stream_event(
                    {
                        "event": "error",
                        "ok": False,
                        "mode": "conversation",
                        "error": str(fallback_exc),
                        "elapsed_ms": int((time.time() - started) * 1000),
                    }
                )
        self.send_stream_event({"event": "done"})

    def handle_structured_stream(self, body: dict[str, Any], started: float) -> None:
        task = str(body.get("task") or "").strip()
        company = str(body.get("company") or "").strip() or None
        if not task:
            self.send_json({"ok": False, "error": "请输入招商任务"}, 400)
            return

        self.start_ndjson_stream()
        self.send_stream_event({"event": "stage", "status": "active", "label": "解析目标", "detail": "识别招商目标、数量要求与交付方向", "elapsed_ms": 0})

        if detect_intent(task) == "data_inventory":
            stats = data_stats()
            self.send_stream_event({"event": "stage", "status": "done", "label": "整理资料覆盖", "detail": f"企业 {stats.get('enterprise_count')} 家，政策 {stats.get('policy_count')} 条"})
            self.send_stream_event({"event": "stats", "ok": True, "intent": "data_inventory", "stats": stats, "elapsed_ms": int((time.time() - started) * 1000)})
            self.send_stream_event({"event": "done"})
            return

        context = collect_context(task, company)
        if not context.get("ok"):
            self.send_stream_event({"event": "error", "ok": False, **context})
            self.send_stream_event({"event": "done"})
            return

        for step in context.get("steps", [])[:5]:
            self.send_stream_event({"event": "stage", **step, "elapsed_ms": int((time.time() - started) * 1000)})
            time.sleep(0.12)
        self.send_stream_event(
            {
                "event": "context",
                "intent": context.get("intent"),
                "context": compact_context_for_llm(context),
                "sources": context.get("sources", []),
                "evidence": context.get("evidence", []),
                "steps": context.get("steps", []),
                "actions": context.get("actions", []),
            }
        )
        self.send_stream_event({"event": "status", "label": "正在生成招商建议", "detail": "正在综合企业画像、政策匹配与风险信息生成招商研判..."})
        chunks: list[str] = []

        def on_chunk(piece: str) -> None:
            chunks.append(piece)
            self.send_stream_event({"event": "text_delta", "content": piece})

        try:
            report = generate_report_stream(context, on_chunk=on_chunk)
            self.send_stream_event(
                {
                    "event": "artifact",
                    "type": "recommendation",
                    "ok": True,
                    "elapsed_ms": int((time.time() - started) * 1000),
                    "intent": context.get("intent"),
                    "context": compact_context_for_llm(context),
                    "sources": context.get("sources", []),
                    "evidence": context.get("evidence", []),
                    "steps": context.get("steps", []),
                    "actions": context.get("actions", []),
                    "report": report,
                }
            )
        except Exception as exc:
            self.send_stream_event(
                {
                    "event": "artifact",
                    "type": "recommendation",
                    "ok": True,
                    "draft": True,
                    "elapsed_ms": int((time.time() - started) * 1000),
                    "intent": context.get("intent"),
                    "message": "模型暂时不可用，已展示 Agent 完成的候选筛选与依据底稿。",
                    "context": compact_context_for_llm(context),
                    "sources": context.get("sources", []),
                    "evidence": context.get("evidence", []),
                    "steps": context.get("steps", []),
                    "actions": context.get("actions", []),
                    "report": draft_report(context, "模型暂时不可用，已展示 Agent 完成的候选筛选与依据底稿。"),
                }
            )
        self.send_stream_event({"event": "done"})

    def do_POST(self) -> None:
        if self.path == "/api/message_stream":
            started = time.time()
            body = self.read_body()
            task = str(body.get("task") or "").strip()
            if not task:
                self.send_json({"ok": False, "error": "请输入内容"}, 400)
                return

            contextualized_task = contextualize_task_for_backend(task, body)
            if should_use_original_agent_chat(task) or not should_start_structured_workflow(task, body):
                self.handle_original_agent_stream(contextualized_task, started)
            else:
                body["task"] = contextualized_task
                self.handle_structured_stream(body, started)
            return

        if self.path == "/api/analyze_stream":
            started = time.time()
            body = self.read_body()
            raw_task = str(body.get("task") or "").strip()
            body["task"] = contextualize_task_for_backend(raw_task, body)
            self.handle_structured_stream(body, started)
            return

        if self.path == "/api/analyze":
            started = time.time()
            body = self.read_body()
            task = str(body.get("task") or "").strip()
            company = str(body.get("company") or "").strip() or None
            if not task:
                self.send_json({"ok": False, "error": "请输入招商任务"}, 400)
                return

            if detect_intent(task) == "data_inventory":
                self.send_json(
                    {
                        "ok": True,
                        "elapsed_ms": int((time.time() - started) * 1000),
                        "intent": "data_inventory",
                        "stats": data_stats(),
                    }
                )
                return

            context = collect_context(contextualize_task_for_backend(task, body), company)
            if not context.get("ok"):
                self.send_json({"ok": False, **context}, 404)
                return
            try:
                report = generate_report(context)
            except Exception as exc:
                draft = draft_report(context, "模型暂时不可用，已展示 Agent 完成的候选筛选与依据底稿。")
                self.send_json(
                    {
                        "ok": True,
                        "draft": True,
                        "error": str(exc),
                        "message": "模型暂时不可用，已展示 Agent 完成的候选筛选与依据底稿。",
                        "intent": context.get("intent"),
                        "context": compact_context_for_llm(context),
                        "sources": context.get("sources", []),
                        "evidence": context.get("evidence", []),
                        "steps": context.get("steps", []),
                        "actions": context.get("actions", []),
                        "report": draft,
                    },
                )
                return
            self.send_json(
                {
                    "ok": True,
                    "elapsed_ms": int((time.time() - started) * 1000),
                    "intent": context.get("intent"),
                    "context": compact_context_for_llm(context),
                    "sources": context.get("sources", []),
                    "evidence": context.get("evidence", []),
                    "steps": context.get("steps", []),
                    "actions": context.get("actions", []),
                    "report": report,
                }
            )
            return

        if self.path == "/api/material":
            body = self.read_body()
            task = str(body.get("task") or "").strip()
            company = str(body.get("company") or "").strip() or None
            material_type = str(body.get("type") or "outline").strip()
            report = body.get("report") or {}
            context = collect_context(task, company)
            if not context.get("ok"):
                self.send_json({"ok": False, **context}, 404)
                return
            try:
                material = generate_material(material_type, report, context)
            except Exception as exc:
                self.send_json(
                    {
                        "ok": False,
                        "error": str(exc),
                        "message": "LLM 不可用，已停止生成材料；不会使用固定话术冒充结果。",
                        "sources": context.get("sources", []),
                        "evidence": context.get("evidence", []),
                    },
                    502,
                )
                return
            self.send_json(
                {
                    "ok": True,
                    "material": material,
                    "sources": context.get("sources", []),
                    "evidence": context.get("evidence", []),
                }
            )
            return

        if self.path == "/api/export/report":
            body = self.read_body()
            report = body.get("report") or {}
            context = body.get("context") or {}
            export_format = str(body.get("format") or "docx").strip().lower()
            if not report:
                self.send_json({"ok": False, "error": "缺少报告数据"}, 400)
                return
            try:
                if export_format == "pdf":
                    out_bytes = build_report_pdf(report, context)
                    mime = "application/pdf"
                else:
                    out_bytes = build_report_docx(report, context)
                    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            except Exception as exc:
                self.send_json({"ok": False, "error": f"报告导出失败：{exc}"}, 500)
                return
            ext = "pdf" if export_format == "pdf" else "docx"
            date_str = time.strftime("%Y%m%d")
            raw = f"园区重点招商企业推荐建议_{date_str}.{ext}"
            filename = urllib.parse.quote(raw)
            self.send_response(200)
            self.send_header("Content-Type", mime)
            self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{filename}")
            self.send_header("Content-Length", str(len(out_bytes)))
            self.send_header("Cache-Control", "no-store")
            self.send_cors_headers()
            self.end_headers()
            self.wfile.write(out_bytes)
            return

        if self.path == "/api/export/material":
            body = self.read_body()
            material = body.get("material") or {}
            material_type = str(body.get("type") or "outline").strip()
            export_format = str(body.get("format") or "docx").strip().lower()
            if not material:
                self.send_json({"ok": False, "error": "缺少材料数据"}, 400)
                return
            try:
                if export_format == "pdf":
                    out_bytes = build_material_pdf(material, material_type)
                    mime = "application/pdf"
                else:
                    out_bytes = build_material_docx(material, material_type)
                    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            except Exception as exc:
                self.send_json({"ok": False, "error": f"材料导出失败：{exc}"}, 500)
                return
            ext = "pdf" if export_format == "pdf" else "docx"
            date_str = time.strftime("%Y%m%d")
            company = _safe_text(material.get("company_name") or "")
            if material_type == "outline":
                raw = f"关于赴{company}开展招商拜访的提纲_{date_str}.{ext}" if company else f"招商拜访提纲_{date_str}.{ext}"
            elif material_type == "wechat":
                raw = f"{company}招商跟进话术_{date_str}.{ext}" if company else f"招商跟进话术_{date_str}.{ext}"
            else:
                label = MATERIAL_LABELS.get(material_type, material_type or "招商材料")
                raw = f"{label}_{date_str}.{ext}"
            filename = urllib.parse.quote(raw)
            self.send_response(200)
            self.send_header("Content-Type", mime)
            self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{filename}")
            self.send_header("Content-Length", str(len(out_bytes)))
            self.send_header("Cache-Control", "no-store")
            self.send_cors_headers()
            self.end_headers()
            self.wfile.write(out_bytes)
            return

        self.send_json({"ok": False, "error": "未知接口"}, 404)

    def serve_static(self) -> None:
        raw_path = self.path.split("?", 1)[0].split("#", 1)[0]
        if raw_path in ("", "/"):
            raw_path = "/index.html"
        safe = Path(raw_path.lstrip("/"))
        file_path = (STATIC_DIR / safe).resolve()
        if not str(file_path).startswith(str(STATIC_DIR.resolve())) or not file_path.exists() or file_path.is_dir():
            self.send_error(404)
            return
        content = file_path.read_bytes()
        mime = mimetypes.guess_type(str(file_path))[0] or "application/octet-stream"
        if file_path.suffix == ".js":
            mime = "text/javascript"
        self.send_response(200)
        self.send_header("Content-Type", f"{mime}; charset=utf-8" if mime.startswith("text/") else mime)
        self.send_header("Content-Length", str(len(content)))
        self.end_headers()
        self.wfile.write(content)


def _rag_startup_check() -> str:
    """Quick RAG health check at startup so operators can see status in console."""
    try:
        results = rag_for("health check", None, "")
        if results and not results[0].get("unavailable"):
            return f"OK ({len(results)} 条片段)"
        reason = (results[0].get("metadata") or {}).get("reason", "unknown") if results else "empty"
        error = (results[0].get("metadata") or {}).get("error", "")[:80] if results else ""
        return f"不可用 (reason={reason}, {error})"
    except Exception as exc:
        return f"异常 ({exc})"


def main() -> None:
    port = int(os.getenv("MVP_PORT", "8765"))
    host = os.getenv("MVP_HOST", "127.0.0.1")
    server = ThreadingHTTPServer((host, port), Handler)
    print(f"ParkFlow AI MVP running at http://{host}:{port}")
    print(f"Reading original V0 data from: {ORIGINAL_DIR}")
    rag_status = _rag_startup_check()
    print(f"RAG / ChromaDB: {rag_status}")
    print("Press Ctrl+C to stop.")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        pass
    finally:
        server.server_close()


if __name__ == "__main__":
    main()
